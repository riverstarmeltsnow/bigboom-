"""
基于实验二指导书 + 实验四模板格式生成报告
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(BASE, "images")
OUTPUT = os.path.join(BASE, "output", "2024121065-何鑫-网工242-密码学实验2.docx")
CODE_FILE = os.path.join(BASE, "templates", "_code.txt")

# 截图（当前为实验2占位，用户需替换）
SCREENSHOTS = []

# 内容行: (text, size, bold, font_name, line_sp, align)
CONTENT = [
    # 标题
    ("应用密码学实验报告", 16, True, "宋体", None, "center"),
    ("课程：应用密码学    实验名称：RSA 算法的实现", 14, False, "宋体", None, None),
    ("姓名：何鑫            实验日期：2026.5.8", 14, False, "宋体", None, None),
    ("学号：2024121065      实验报告日期：2026.5.8", 14, False, "宋体", None, None),
    ("班级：网工242", 14, False, "宋体", None, None),
    # 一
    ("一、实验内容", 14, True, "宋体", None, None),
    ("RSA 算法实现实验：使用平方乘算法和模重复平方法，实现 RSA 加密和解密功能，使用扩展欧几里得算法计算私钥。", 12, False, "宋体", 1.3, None),
    # 二
    ("二、实验环境（详细说明所用的系统及平台）", 14, True, "宋体", None, None),
    ("操作系统：Windows 11", 12, False, "Times New Roman", 1.3, None),
    ("开发平台：Visual Studio 2022", 12, False, "Times New Roman", 1.3, None),
    ("编程语言：C++", 12, False, "Times New Roman", 1.3, None),
    # 三
    ("三、实验目的", 14, True, "宋体", None, None),
    ("（1）加深对 RSA 算法的理解；", 12, False, "宋体", 1.3, None),
    ("（2）加深对平方乘算法和模重复平方法的理解；", 12, False, "宋体", 1.3, None),
    ("（3）加深对模块化设计的理解，提高编程实践能力。", 12, False, "宋体", 1.3, None),
    # 四
    ("四、实验内容、步骤及结果", 14, True, "宋体", None, None),
    ("1. 实验内容", 12, True, "宋体", 1.5, None),
    ("（1）完成 RSA 算法加密和解密功能。", 12, False, "宋体", 1.3, None),
    ("（2）按照欧几里得扩展算法求，计算 RSA 私钥。", 12, False, "宋体", 1.3, None),
    ("（3）按照平方乘算法和模重复平方法，分别计算 a^m mod n，完成 RSA 的加密和解密。", 12, False, "宋体", 1.3, None),
    ("2. 实验步骤", 12, True, "宋体", 1.5, None),
    ("步骤1：确定参数。取素数 p=13, q=17，计算 n=p×q=221，φ(n)=(p-1)(q-1)=192，公钥 e=7。", 12, False, "Times New Roman", 1.3, None),
    ("步骤2：使用扩展欧几里得算法计算私钥 d：d = e⁻¹ mod φ(n) = 7⁻¹ mod 192。", 12, False, "Times New Roman", 1.3, None),
    ("步骤3：第一组测试，固定消息 m=22。将 m 转换为二进制，调用平方乘算法 LRFun_65 加密得密文 c，再调用 RLFun_65 解密验证。", 12, False, "Times New Roman", 1.3, None),
    ("步骤4：第二组测试，学号尾号为 5，取 m=15。重复加密和解密过程，输出中间结果。", 12, False, "Times New Roman", 1.3, None),
    ("步骤5：输出中间计算结果，比较加密解密结果是否一致。", 12, False, "Times New Roman", 1.3, None),
    # 五
    ("五、实验结果与分析", 14, True, "宋体", None, None),
    ("【图片区域】", 0, False, None, None, None),  # 占位
    ("结果分析：程序编译运行正常。平方乘算法和模重复平方法均能正确计算 a^m mod n，加密结果和解密后恢复的明文一致，RSA 算法实现正确。两组测试均验证通过。", 12, False, "宋体", 1.3, None),
    # 六
    ("六、实验结论", 14, True, "宋体", None, None),
    ("本次实验通过 C++ 编程实现了 RSA 加密解密算法，使用平方乘算法和模重复平方法完成了 RSA 的加密和解密计算，使用扩展欧几里得算法求出了私钥 d。实验结果表明，RSA 算法能够正确实现消息的加密与解密，达到了实验的预期目标。", 12, False, "宋体", 1.3, None),
    # 七
    ("七、附录（核心代码）", 14, True, "宋体", None, None),
    ("【代码区域】", 0, False, None, None, None),
]


def set_run_font(run, font_name, size, bold):
    """设置 run 字体"""
    run.font.name = font_name if font_name else "宋体"
    run.font.size = Pt(size)
    run.font.bold = bold
    if font_name:
        run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def set_para_format(p, line_sp=None, align=None):
    """设置段落格式"""
    if line_sp is not None:
        p.paragraph_format.line_spacing = line_sp
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)


def clear_para(p):
    """清空段落所有 run"""
    for run in p.runs:
        run.text = ""


def generate():
    doc = Document()
    # ===== 页面设置 =====
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    ps = doc.paragraphs  # initially len=1

    # 写内容
    content_idx = 0
    for item in CONTENT:
        text, size, bold, font_name, line_sp, align = item

        if text == "【图片区域】":
            if SCREENSHOTS:
                for img_name, caption in SCREENSHOTS:
                    img_path = os.path.join(IMG_DIR, img_name)
                    if os.path.exists(img_path):
                        # 新段落放图片
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(img_path, width=Inches(5.0))
                        content_idx += 1

                        # 图注
                        cp = doc.add_paragraph()
                        run = cp.add_run(caption)
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cp.paragraph_format.space_after = Pt(4)
                        content_idx += 1
                    else:
                        p = doc.add_paragraph()
                        run = p.add_run(f"【图片缺失: {img_name}】")
                        run.font.size = Pt(10)
                        content_idx += 1
            else:
                p = doc.add_paragraph()
                run = p.add_run("【请在此处插入运行结果截图】")
                set_run_font(run, "宋体", 10, False)
                content_idx += 1

        elif text == "【代码区域】":
            if os.path.exists(CODE_FILE):
                with open(CODE_FILE, "r", encoding="utf-8") as f:
                    code_text = f.read()
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(7.5)
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(4)
                content_idx += 1
            else:
                p = doc.add_paragraph()
                run = p.add_run("【代码文件缺失】")
                content_idx += 1

        else:
            p = doc.add_paragraph()
            run = p.add_run(text)
            set_run_font(run, font_name, size, bold)
            set_para_format(p, line_sp, align)
            content_idx += 1

    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    print(f"[OK] 实验报告已保存: {OUTPUT}")
    print(f"  段落数: {len(doc.paragraphs)}, 内容行: {content_idx}")


if __name__ == "__main__":
    generate()
