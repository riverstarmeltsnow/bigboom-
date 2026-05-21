"""
生成实验报告：Session 实现用户登录与注销 + Cookie 记录上次访问时间
仿照 2024121065-何鑫-实验2.docx 格式
"""
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(BASE, "output", "2024121065-何鑫-SessionCookie实验.docx")
CODE_FILE = os.path.join(BASE, "templates", "_code_session_cookie.txt")

# ========== 内容 ==========
CONTENT = [
    # 信息头 (16pt)
    ("课程名称：  Web 应用开发", 16, False),
    ("班    级：  网工242", 16, False),
    ("学生姓名：  何鑫", 16, False),
    ("学    号：  2024121065", 16, False),
    ("实验日期：  2026/5/14", 16, False),
    ("", 0, False),  # 空行

    # 一、实验目的
    ("实验目的", 18, False),
    ("实验 1：掌握 Session 的工作原理，实现用户登录与注销功能。学习如何在 Servlet 中创建和销毁 Session，在 JSP 中读取 Session 数据，以及通过 Session 实现受保护页面的访问控制。",
     12, False),
    ("实验 2：掌握 Cookie 的读写操作，实现记录用户上次访问时间的功能。学习如何创建 Cookie、设置存活期、读取 Cookie 数据，以及处理首次访问和 Cookie 值编码。",
     12, False),

    # 二、实验内容与步骤
    ("实验内容与步骤", 18, False),
    ("实验 1：Session 实现用户登录与注销", 12, True),
    ("（1）创建登录页面 login.jsp，包含用户名和密码输入框，提交到 LoginServlet。",
     12, False),
    ("（2）LoginServlet 硬编码合法用户 hexin/123456。成功则将用户名存入 Session 并重定向到 welcome.jsp；失败则重定向回 login.jsp 并显示错误信息。",
     12, False),
    ("（3）welcome.jsp 从 Session 获取用户名显示欢迎信息，包含「退出登录」链接和「个人中心」链接。",
     12, False),
    ("（4）LogoutServlet 调用 session.invalidate() 销毁会话后重定向到 login.jsp。",
     12, False),
    ("（5）profile.jsp 为受保护页面，未登录用户直接访问时自动跳转到 login.jsp 并提示「请先登录」。",
     12, False),
    ("实验 2：Cookie 记录上次访问时间", 12, True),
    ("（1）用户成功登录后，welcome.jsp 读取名为 lastVisit 的 Cookie，显示用户上次访问时间。",
     12, False),
    ("（2）如果 Cookie 不存在，显示「这是您第一次登录」。", 12, False),
    ("（3）每次用户访问 welcome.jsp 时，更新 Cookie：将当前时间写入 lastVisit，存活期设为 7 天。",
     12, False),
    ("（4）Cookie 值使用 URLEncoder/URLDecoder 编码解码，避免 RFC 6265 非法字符错误。", 12, False),

    # 三、核心代码
    ("核心代码", 18, False),
]

# ========== 页面设置 ==========
def setup_page(doc):
    section = doc.sections[0]
    section.page_width = Pt(595.3)
    section.page_height = Pt(841.9)
    section.left_margin = Pt(90)
    section.right_margin = Pt(90)
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)


def add_para(doc, text, size, bold, font_name="宋体", align=None, line_sp=None):
    """添加段落"""
    p = doc.add_paragraph()
    if not text and size == 0:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(10)
        return p

    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    if line_sp is not None:
        p.paragraph_format.line_spacing = line_sp
    return p


def add_code_para(doc, text, size=10.5, font_name="Consolas"):
    """添加代码段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    return p


def generate():
    doc = Document()
    setup_page(doc)

    # 写入前置内容
    for text, size, bold in CONTENT:
        add_para(doc, text, size, bold)

    # 读取代码文件
    if os.path.exists(CODE_FILE):
        with open(CODE_FILE, "r", encoding="utf-8") as f:
            code_text = f.read()

        # 按文件名分割
        sections = code_text.strip().split("=====")
        for section in sections:
            section = section.strip()
            if not section:
                continue
            lines = section.split("\n", 1)
            filename = lines[0].strip().rstrip("=").strip()
            if len(lines) > 1:
                add_code_para(doc, filename, size=10.5, font_name="宋体")
                # 代码内容
                code_body = lines[1].strip()
                # 分行写入
                for line in code_body.split("\n"):
                    add_code_para(doc, line)
                add_code_para(doc, "")  # 空行分隔
    else:
        add_para(doc, "【代码文件缺失】", 12, False)

    # 实验运行截图
    add_para(doc, "实验运行截图", 18, False)
    add_para(doc, "登录页面（含错误提示）：", 12, True)
    add_para(doc, "【请在此处插入 login.jsp 运行截图，显示错误提示信息】", 12, False)
    add_para(doc, "欢迎页面（显示上次访问时间）：", 12, True)
    add_para(doc, "【请在此处插入 welcome.jsp 运行截图，显示欢迎信息与上次访问时间】", 12, False)
    add_para(doc, "个人中心页面（受保护页面）：", 12, True)
    add_para(doc, "【请在此处插入 profile.jsp 运行截图，显示登录用户信息】", 12, False)
    add_para(doc, "未登录直接访问个人中心被拦截：", 12, True)
    add_para(doc, "【请在此处插入未登录访问 profile.jsp 时跳转回 login.jsp 并提示的截图】", 12, False)
    add_para(doc, "第一次登录效果（无 Cookie）：", 12, True)
    add_para(doc, "【请在此处插入首次登录 welcome.jsp 显示「这是您第一次登录」的截图】", 12, False)

    # 实验总结
    add_para(doc, "实验总结", 18, False)
    add_para(doc, "问题 1：使用 response.sendRedirect() 跳转时，request 属性无法传递到新页面，导致错误信息丢失。"
             "解决方法：改用 request.getRequestDispatcher().forward() 传递属性，或使用 URL 参数传递消息。", 12, False)
    add_para(doc, "问题 2：Tomcat 9 对 Cookie 值有严格的 RFC 6265 字符校验，Date().toLocaleString() 中的空格"
             "会触发 IllegalArgumentException。解决方法：对 Cookie 值进行 URL 编码。", 12, False)
    add_para(doc, "问题 3：JSP 中的 form action 使用相对路径时，在嵌套路径下可能出现提交地址错误。"
             "解决方法：使用 ${pageContext.request.contextPath}/path 确保路径正确。", 12, False)

    # 最终加粗总结
    add_para(doc, "实验总结：本次实验实现了基于 Session 的用户登录与注销功能，以及基于 Cookie 的上次访问时间记录功能。"
             "通过实验掌握了 Session 的创建、读取、销毁操作，Cookie 的读写与编码处理，以及受保护页面的访问控制方法。"
             "对 Java Web 开发中的会话管理有了深入理解。", 16, True)

    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    print(f"[OK] 实验报告已保存: {OUTPUT}")
    print(f"  段落数: {len(doc.paragraphs)}")


if __name__ == "__main__":
    generate()
